from __future__ import annotations

import re
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Iterable

import pandas as pd
from docx import Document as DocxDocument
from PIL import Image
from pypdf import PdfReader

from .models import Chunk, Document


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv", ".xlsx", *IMAGE_EXTENSIONS}


def normalize_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text or "")
    return text.strip()


def load_documents_from_paths(paths: Iterable[Path]) -> list[Document]:
    documents: list[Document] = []
    for path in paths:
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS or not path.exists():
            continue
        text = _read_path(path)
        if text:
            documents.append(
                Document(
                    source=path.name,
                    text=normalize_text(text),
                    metadata={"path": str(path), "extension": path.suffix.lower()},
                )
            )
    return documents


def load_uploaded_files(uploaded_files: Iterable[object]) -> list[Document]:
    documents: list[Document] = []
    for uploaded_file in uploaded_files:
        name = getattr(uploaded_file, "name", "uploaded.txt")
        suffix = Path(name).suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            continue
        with NamedTemporaryFile(delete=False, suffix=suffix) as temp:
            temp.write(uploaded_file.getvalue())
            temp_path = Path(temp.name)
        try:
            text = _read_path(temp_path)
            if text:
                documents.append(
                    Document(
                        source=name,
                        text=normalize_text(text),
                        metadata={"extension": suffix, "origin": "upload"},
                    )
                )
        finally:
            temp_path.unlink(missing_ok=True)
    return documents


def chunk_documents(
    documents: Iterable[Document],
    chunk_size: int = 850,
    overlap: int = 140,
) -> list[Chunk]:
    chunks: list[Chunk] = []
    for doc_idx, document in enumerate(documents):
        words = document.text.split()
        if not words:
            continue
        step = max(1, chunk_size - overlap)
        for chunk_idx, start in enumerate(range(0, len(words), step)):
            selected = words[start : start + chunk_size]
            if not selected:
                continue
            chunks.append(
                Chunk(
                    id=f"d{doc_idx}_c{chunk_idx}",
                    source=document.source,
                    text=" ".join(selected),
                    metadata={**document.metadata, "word_start": start},
                )
            )
    return chunks


def _read_path(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        document = DocxDocument(str(path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    if suffix == ".csv":
        frame = pd.read_csv(path)
        return _frame_to_text(frame, source=path.name)
    if suffix == ".xlsx":
        frames = pd.read_excel(path, sheet_name=None)
        return "\n".join(
            _frame_to_text(frame, source=f"{path.name}:{sheet}") for sheet, frame in frames.items()
        )
    if suffix in IMAGE_EXTENSIONS:
        return _image_to_text(path)
    return ""


def _frame_to_text(frame: pd.DataFrame, source: str) -> str:
    records = []
    for idx, row in frame.fillna("").iterrows():
        values = "; ".join(f"{column}: {row[column]}" for column in frame.columns)
        records.append(f"{source} row {idx + 1}: {values}")
    return "\n".join(records)


def _image_to_text(path: Path) -> str:
    with Image.open(path) as image:
        width, height = image.size
        mode = image.mode
    parent = path.parent.name.lower()
    if "схем" in parent or "scheme" in parent:
        kind = "схема флотации"
    elif "регламент" in parent or "equipment" in parent:
        kind = "регламентный визуальный материал"
    else:
        kind = "загруженное изображение"
    return (
        f"{kind}: {path.name}. Формат {path.suffix.lower()}, размер {width}x{height}, режим {mode}. "
        "Учитывать как визуальный источник ограничений по оборудованию, схеме процесса или регламенту. "
        "Если на изображении есть подписи, пользователь должен продублировать критичные параметры в текстовом описании."
    )
