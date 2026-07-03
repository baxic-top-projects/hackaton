from __future__ import annotations

import json
from io import BytesIO
from dataclasses import asdict
from pathlib import Path

import pandas as pd
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from .models import Hypothesis, ResearchBrief


PDF_FONT = "DejaVuSans"


def hypotheses_to_frame(hypotheses: list[Hypothesis]) -> pd.DataFrame:
    rows = []
    for idx, hypothesis in enumerate(hypotheses, start=1):
        rows.append(
            {
                "rank": idx,
                "title": hypothesis.title,
                "total_score": hypothesis.total_score,
                "novelty": hypothesis.novelty,
                "feasibility": hypothesis.feasibility,
                "expected_value": hypothesis.expected_value,
                "risk": hypothesis.risk,
                "confidence": hypothesis.confidence,
                "sources": ", ".join(sorted({item.source for item in hypothesis.evidence})),
                "calculators": "; ".join(
                    f"{item.name}: {item.status} ({item.value})" for item in hypothesis.calculations
                ),
                "statement": hypothesis.statement,
            }
        )
    return pd.DataFrame(rows)


def hypotheses_to_json(hypotheses: list[Hypothesis], brief: ResearchBrief) -> str:
    payload = {
        "brief": asdict(brief),
        "hypotheses": [asdict(hypothesis) for hypothesis in hypotheses],
    }
    return json.dumps(payload, ensure_ascii=False, indent=2)


def hypotheses_to_markdown(hypotheses: list[Hypothesis], brief: ResearchBrief) -> str:
    lines = [
        "# Отчет: Фабрика гипотез",
        "",
        f"**Цель:** {brief.target}",
        f"**Ограничения:** {brief.constraints or 'не указаны'}",
        "",
        "## Ранжированные гипотезы",
        "",
    ]
    for idx, hypothesis in enumerate(hypotheses, start=1):
        lines.extend(
            [
                f"### {idx}. {hypothesis.title}",
                "",
                f"**Итоговый балл:** {hypothesis.total_score:.3f}",
                "",
                hypothesis.statement,
                "",
                f"**Механизм:** {hypothesis.mechanism}.",
                "",
                f"**Обоснование:** {hypothesis.rationale}",
                "",
                "**Расчетные проверки:**",
                *[
                    f"- {item.name}: {item.status}, {item.value}. {item.rationale}"
                    for item in hypothesis.calculations
                ],
                "",
                "**Риски:**",
                *[f"- {risk}" for risk in hypothesis.risks],
                "",
                "**План проверки:**",
                *[f"- {step}" for step in hypothesis.experiment_plan],
                "",
                "**Источники:**",
                *[f"- {item.source}: {item.quote}" for item in hypothesis.evidence],
                "",
            ]
        )
    return "\n".join(lines)


def hypotheses_to_docx(hypotheses: list[Hypothesis], brief: ResearchBrief) -> bytes:
    document = Document()
    document.add_heading("Отчет: Фабрика гипотез", level=1)
    document.add_paragraph(f"Цель: {brief.target}")
    document.add_paragraph(f"Ограничения: {brief.constraints or 'не указаны'}")
    document.add_paragraph(f"Доступные материалы: {brief.available_materials or 'не указаны'}")
    document.add_paragraph(f"Оборудование: {brief.equipment or 'не указано'}")

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["#", "Гипотеза", "Score", "Новизна", "Риск", "Источники"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for rank, hypothesis in enumerate(hypotheses, start=1):
        row = table.add_row().cells
        row[0].text = str(rank)
        row[1].text = hypothesis.title
        row[2].text = f"{hypothesis.total_score:.3f}"
        row[3].text = f"{hypothesis.novelty:.3f}"
        row[4].text = f"{hypothesis.risk:.3f}"
        row[5].text = ", ".join(sorted({item.source for item in hypothesis.evidence}))

    for rank, hypothesis in enumerate(hypotheses, start=1):
        document.add_heading(f"{rank}. {hypothesis.title}", level=2)
        document.add_paragraph(hypothesis.statement)
        document.add_paragraph(f"Механизм: {hypothesis.mechanism}.")
        document.add_paragraph(f"Обоснование: {hypothesis.rationale}")

        document.add_heading("Расчетные проверки", level=3)
        for item in hypothesis.calculations:
            document.add_paragraph(
                f"{item.name}: {item.status}, {item.value}. {item.rationale}",
                style="List Bullet",
            )

        document.add_heading("Риски", level=3)
        for risk in hypothesis.risks:
            document.add_paragraph(risk, style="List Bullet")

        document.add_heading("План проверки", level=3)
        for step in hypothesis.experiment_plan:
            document.add_paragraph(step, style="List Number")

        document.add_heading("Источники", level=3)
        for evidence in hypothesis.evidence:
            document.add_paragraph(
                f"{evidence.source} (релевантность {evidence.score:.3f}): {evidence.quote}",
                style="List Bullet",
            )

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def hypotheses_to_pdf(hypotheses: list[Hypothesis], brief: ResearchBrief) -> bytes:
    font_name = _register_pdf_font()
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=1.5 * cm,
        leftMargin=1.5 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
    )
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "BaseCyrillic",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=9,
        leading=12,
    )
    heading = ParagraphStyle(
        "HeadingCyrillic",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=16,
        leading=20,
    )
    subheading = ParagraphStyle(
        "SubheadingCyrillic",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=12,
        leading=15,
        spaceBefore=10,
    )
    story = [
        Paragraph("Отчет: Фабрика гипотез", heading),
        Paragraph(f"<b>Цель:</b> {_escape(brief.target)}", base),
        Paragraph(f"<b>Ограничения:</b> {_escape(brief.constraints or 'не указаны')}", base),
        Spacer(1, 0.3 * cm),
    ]

    rows = [["#", "Гипотеза", "Score", "Новизна", "Риск"]]
    for idx, hypothesis in enumerate(hypotheses, start=1):
        rows.append(
            [
                str(idx),
                Paragraph(_escape(hypothesis.title), base),
                f"{hypothesis.total_score:.3f}",
                f"{hypothesis.novelty:.3f}",
                f"{hypothesis.risk:.3f}",
            ]
        )
    table = Table(rows, colWidths=[0.8 * cm, 10 * cm, 1.8 * cm, 2 * cm, 1.6 * cm])
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.extend([table, Spacer(1, 0.4 * cm)])

    for idx, hypothesis in enumerate(hypotheses, start=1):
        story.append(Paragraph(f"{idx}. {_escape(hypothesis.title)}", subheading))
        story.append(Paragraph(_escape(hypothesis.statement), base))
        story.append(Paragraph(f"<b>Механизм:</b> {_escape(hypothesis.mechanism)}.", base))
        story.append(Paragraph(f"<b>Обоснование:</b> {_escape(hypothesis.rationale)}", base))
        story.append(Paragraph("<b>Расчетные проверки:</b>", base))
        for item in hypothesis.calculations:
            story.append(Paragraph(f"- {_escape(item.name)}: {item.status}, {_escape(item.value)}", base))
        story.append(Paragraph("<b>План проверки:</b>", base))
        for step in hypothesis.experiment_plan:
            story.append(Paragraph(f"- {_escape(step)}", base))
        story.append(Paragraph("<b>Источники:</b>", base))
        for evidence in hypothesis.evidence[:5]:
            story.append(
                Paragraph(
                    f"- {_escape(evidence.source)} ({evidence.score:.3f}): {_escape(evidence.quote[:350])}",
                    base,
                )
            )
        story.append(Spacer(1, 0.25 * cm))

    document.build(story)
    return output.getvalue()


def _register_pdf_font() -> str:
    if PDF_FONT in pdfmetrics.getRegisteredFontNames():
        return PDF_FONT
    candidates = [
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/DejaVuSans.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            pdfmetrics.registerFont(TTFont(PDF_FONT, str(candidate)))
            return PDF_FONT
    return "Helvetica"


def _escape(text: str) -> str:
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
